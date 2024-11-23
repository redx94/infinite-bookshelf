# 1: Import libraries
import streamlit as st
from groq import Groq
import json
import os
import io
import ebooklib
from ebooklib import epub
from docx import Document

from infinite_bookshelf.agents import (
    generate_section,
    generate_book_structure,
    generate_book_title,
)
from infinite_bookshelf.inference import GenerationStatistics
from infinite_bookshelf.tools import create_markdown_file, create_pdf_file
from infinite_bookshelf.ui.components import (
    render_groq_form,
    display_statistics,
)
from infinite_bookshelf.ui import Book, load_return_env, ensure_states

# Utility Functions for Book Export
def create_epub_file(book):
    """Create an ePub file from the book content"""
    epub_book = epub.EpubBook()
    epub_book.set_title(book.book_title)
    epub_book.set_language('en')

    # Create chapters
    chapters = []
    for title, content in book.contents.items():
        chapter = epub.EpubHtml(title=title, file_name=f'{title}.xhtml', lang='en')
        chapter.set_content(f'<h1>{title}</h1><p>{content}</p>')
        epub_book.add_item(chapter)
        chapters.append(chapter)

    # Create Table of Contents
    epub_book.toc = tuple(chapters)
    epub_book.add_item(epub.EpubNcx())
    epub_book.add_item(epub.EpubNav())

    # Basic styling
    style = 'BODY {color: white;}'
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
    epub_book.add_item(nav_css)

    # Write ePub
    epub_buffer = io.BytesIO()
    epub.write_epub(epub_buffer, epub_book, {})
    epub_buffer.seek(0)
    return epub_buffer

def create_docx_file(book):
    """Create a DOCX file from the book content"""
    doc = Document()
    doc.add_heading(book.book_title, level=1)
    
    for title, content in book.contents.items():
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def get_available_groq_models():
    """Fetch available Groq models"""
    try:
        groq = Groq()
        models = groq.models.list()
        return [model.id for model in models.data]
    except Exception as e:
        st.error(f"Could not fetch Groq models: {e}")
        return ["llama3-8b-8192", "llama3-70b-8192", "gemma-7b-it"]

def advanced_book_generation():
    """Advanced book generation workflow"""
    st.title("Advanced Book Generation")
    
    # Model Selection
    available_models = get_available_groq_models()
    
    col1, col2, col3 = st.columns(3)
    with col1:
        title_model = st.selectbox("Title Generation Model", available_models)
    with col2:
        structure_model = st.selectbox("Structure Generation Model", available_models)
    with col3:
        content_model = st.selectbox("Content Generation Model", available_models)
    
    # Book Topic and Instructions
    topic = st.text_input("Book Topic")
    additional_instructions = st.text_area("Additional Instructions")
    
    # Generation Workflow
    if st.button("Generate Book"):
        with st.spinner("Generating Book..."):
            # Title Generation
            book_title = generate_book_title(
                prompt=topic, 
                model=title_model, 
                groq_provider=Groq()
            )
            
            # Structure Generation
            _, book_structure = generate_book_structure(
                prompt=topic,
                additional_instructions=additional_instructions,
                model=structure_model,
                groq_provider=Groq()
            )
            
            # Create Book Object
            book_structure_json = json.loads(book_structure)
            book = Book(book_title, book_structure_json)
            
            # Content Generation
            def generate_book_content(sections):
                for title, content in sections.items():
                    if isinstance(content, str):
                        generated_content = generate_section(
                            prompt=f"{title}: {content}",
                            additional_instructions=additional_instructions,
                            model=content_model,
                            groq_provider=Groq()
                        )
                        book.update_content(title, generated_content)
                    elif isinstance(content, dict):
                        generate_book_content(content)
            
            generate_book_content(book_structure_json)
            
            # Display Book
            st.write(f"# {book.book_title}")
            book.display_structure()
            
            # Export Options
            st.subheader("Export Book")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.download_button(
                    label="Download PDF", 
                    data=create_pdf_file(book.get_markdown_content()),
                    file_name=f"{book_title}.pdf"
                ):
                    st.success("PDF Downloaded")
            
            with col2:
                if st.download_button(
                    label="Download ePub", 
                    data=create_epub_file(book),
                    file_name=f"{book_title}.epub"
                ):
                    st.success("ePub Downloaded")
            
            with col3:
                if st.download_button(
                    label="Download DOCX", 
                    data=create_docx_file(book),
                    file_name=f"{book_title}.docx"
                ):
                    st.success("DOCX Downloaded")
            
            with col4:
                if st.download_button(
                    label="Download Markdown", 
                    data=create_markdown_file(book.get_markdown_content()),
                    file_name=f"{book_title}.md"
                ):
                    st.success("Markdown Downloaded")
            
            # Edit Book Option
            st.subheader("Edit Book")
            edit_mode = st.checkbox("Enable Editing")
            
            if edit_mode:
                for title in book.contents.keys():
                    st.subheader(title)
                    edited_content = st.text_area(
                        f"Edit {title}", 
                        value=book.contents[title]
                    )
                    
                    if st.button(f"Save Changes for {title}"):
                        book.update_content(title, edited_content)
                        st.success(f"Updated {title}")

def main():
    st.sidebar.title("Infinite Bookshelf")
    
    # Navigation
    page = st.sidebar.radio(
        "Choose a Page", 
        ["Advanced Book Generation", "Simple Book Generation"]
    )
    
    if page == "Advanced Book Generation":
        advanced_book_generation()
    else:
        st.write("Simple Book Generation (To be implemented)")

if __name__ == "__main__":
    main()
